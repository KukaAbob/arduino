import json
import re
from docx import Document

def parse_arduino_docx(file_path, debug=True):
    doc = Document(file_path)

    # Шаг 1. Извлекаем вопросы и варианты ответов из основного текста документа.
    question_pattern = re.compile(
        r"Вопрос\s*(\d+)\s*[:.-]?\s*(.*?)\s*A\)\s*(.*?)\s*B\)\s*(.*?)\s*C\)\s*(.*?)\s*D\)\s*(.*?)(?:\s*E\)\s*(.*?))?(?:\n|$)",
        re.MULTILINE | re.DOTALL
    )
    
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    text = "\n".join(paragraphs)
    
    questions = []
    for match in question_pattern.finditer(text):
        try:
            q_id = int(match.group(1))
            question_text = match.group(2).strip()
            options = {
                'A': match.group(3).strip(),
                'B': match.group(4).strip(),
                'C': match.group(5).strip(),
                'D': match.group(6).strip()
            }
            if match.group(7):
                options['E'] = match.group(7).strip()
            
            questions.append({
                'id': q_id,
                'question': question_text,
                'options': options,
                'correct': 'N/A'
            })
        except Exception as e:
            if debug:
                print(f"[DEBUG] Ошибка при обработке вопроса: {str(e)}")
            continue

    # Шаг 2. Извлекаем правильные ответы из таблиц.
    # Предполагаем, что таблицы организованы в виде:
    # В строке: [Вопрос, Ответ, Вопрос, Ответ, Вопрос, Ответ, Вопрос, Ответ]
    answers = {}
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows):
            # Если это заголовочная строка, пропускаем её (например, если первая ячейка равна "Вопрос")
            if row_idx == 0 and "Вопрос" in row.cells[0].text:
                continue
            cells = row.cells
            # Перебираем пары ячеек: (0,1), (2,3), (4,5), (6,7)
            for pair_start in [0, 2, 4, 6]:
                # Если в строке меньше пар, то пропускаем
                if pair_start >= len(cells):
                    break
                q_text = cells[pair_start].text.strip()
                # Если q_text не является числом, пропускаем эту пару
                if not re.match(r"^\d+$", q_text):
                    continue
                q_num = int(q_text)
                # Если следующая ячейка есть, берем из неё ответ
                if pair_start + 1 < len(cells):
                    ans_text = cells[pair_start + 1].text.strip()
                    # Ищем букву (A–E) в ответе
                    ans_match = re.search(r'\b[A-E]\b', ans_text)
                    if ans_match:
                        answers[q_num] = ans_match.group(0)
                        if debug:
                            print(f"[DEBUG] Вопрос {q_num}: Найден правильный ответ '{ans_match.group(0)}' (таблица {table_idx}, строка {row_idx+1}, ячейки {pair_start+1}-{pair_start+2}).")
                    else:
                        if debug:
                            print(f"[DEBUG] Вопрос {q_num}: В ячейке с ответом (ячейка {pair_start+2}) не найдено буквы (текст: '{ans_text}').")
                else:
                    if debug:
                        print(f"[DEBUG] Вопрос {q_num}: Нет ячейки с ответом для пары, начиная с ячейки {pair_start+1}.")

    # Шаг 3. Присваиваем найденные ответы соответствующим вопросам.
    for question in questions:
        if question['id'] in answers:
            question['correct'] = answers[question['id']]
        else:
            if debug:
                print(f"[WARNING] Вопрос {question['id']}: Ответ не найден в таблицах, оставляем 'N/A'.")
    
    return {'questions': sorted(questions, key=lambda x: x['id'])}

if __name__ == '__main__':
    try:
        file_path = "ard.docx"  # Замените на путь к вашему файлу
        data = parse_arduino_docx(file_path, debug=True)
        
        with open('questions_arduino.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        print(f"Успешно обработано {len(data['questions'])} вопросов.")
        print("Файл questions_arduino.json успешно создан.")
    except Exception as e:
        print(f"Ошибка: {str(e)}")
